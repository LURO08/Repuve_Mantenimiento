from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
from tempfile import NamedTemporaryFile

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = ROOT / "templates" / "formatos"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def replace_cell(cell, value):
    cell.text = value


def prepare_checklist(source, destination):
    doc = Document(str(source))
    table = doc.tables[0]
    for index, row in enumerate(table.rows[2:]):
        key = f"CHK_{index}"
        if index == len(table.rows[2:]) - 1:
            replace_cell(row.cells[0], "Otro: {{OTRO_ESPECIFIQUE}}")
        replace_cell(row.cells[1], "{{" + key + "_BUENO}}")
        replace_cell(row.cells[2], "{{" + key + "_MALO}}")
        replace_cell(row.cells[3], "{{" + key + "_OBS}}")
        replace_cell(row.cells[4], "{{" + key + "_CAMBIADO}}")
    doc.save(str(destination))


def prepare_quality(source, destination):
    doc = Document(str(source))
    table = doc.tables[0]
    for index, row in enumerate(table.rows[1:]):
        replace_cell(row.cells[0], "{{CARRIL_" + str(index) + "_NOMBRE}}")
        replace_cell(row.cells[1], "{{CARRIL_" + str(index) + "_LECTURA}}")
        replace_cell(row.cells[2], "{{CARRIL_" + str(index) + "_MONITOREO}}")
        replace_cell(row.cells[3], "{{CARRIL_" + str(index) + "_OBS}}")

    replacements = {
        "Energía funciona correctamente: [ ] Luz Eléctrica (1+1) [ ] Paneles Solares/Baterías":
            "Energía funciona correctamente: {{ENERGIA_LUZ}} Luz Eléctrica (1+1) "
            "{{ENERGIA_SOLAR}} Paneles Solares/Baterías",
        "Enlace funcionando correctamente [ ] Sí [ ] No":
            "Enlace funcionando correctamente {{ENLACE_SI}} Sí {{ENLACE_NO}} No",
        "Sistema de Monitoreo opera al 100%, es accesible y llegan lecturas [ ] Sí [ ] No":
            "Sistema de Monitoreo opera al 100%, es accesible y llegan lecturas "
            "{{MONITOREO_SI}} Sí {{MONITOREO_NO}} No",
        "Resultado de la Prueba: [ ] Exitosa [ ] Fallida":
            "Resultado de la Prueba: {{RESULTADO_EXITOSA}} Exitosa {{RESULTADO_FALLIDA}} Fallida",
        "Acciones Correctivas Realizadas (en caso de falla):":
            "Acciones Correctivas Realizadas (en caso de falla): {{ACCIONES_CORRECTIVAS}}",
    }
    for paragraph in doc.paragraphs:
        normalized = " ".join(paragraph.text.split())
        if normalized in replacements:
            paragraph.text = replacements[normalized]
    doc.save(str(destination))


def prepare_tools(source, destination):
    doc = Document(str(source))
    for table_index, table in enumerate(doc.tables):
        prefix = ("HERR", "CONS", "EPP")[table_index]
        for row_index, row in enumerate(table.rows[1:]):
            replace_cell(row.cells[1], "{{" + prefix + "_" + str(row_index) + "_IZQ}}")
            replace_cell(row.cells[3], "{{" + prefix + "_" + str(row_index) + "_DER}}")
    first_consumable = doc.tables[1].rows[1].cells[0]
    replace_cell(first_consumable, "Antenas Yagi - Cantidad: ({{YAGI_CANTIDAD}})")
    doc.save(str(destination))


def inject_service_fields(path, format_type):
    with ZipFile(path, "r") as source:
        entries = {name: source.read(name) for name in source.namelist()}

    root = etree.fromstring(entries["word/document.xml"])
    for node in root.xpath(".//w:t", namespaces=NS):
        text = node.text or ""
        if "Nombre del Arco:" in text:
            node.text = text.replace("Nombre del Arco:", "Nombre del Arco: {{ARCO}}")
        elif text.strip() == "Fecha:":
            node.text = "Fecha: {{FECHA}}"
        elif "Técnico Responsable" in text or "Técnico Responsable:" in text:
            node.text = text.rstrip(":") + ": {{TECNICO}}"
        elif text.strip() == "Hora:":
            node.text = "Hora: {{HORA}}"
        elif "Ubicación (Referencia):" in text:
            node.text = text.replace("Ubicación (Referencia):", "Ubicación (Referencia): {{UBICACION}}")

        if format_type in {"checklist", "quality"}:
            if text.strip() == "Preventivo":
                node.text = "Preventivo {{TIPO_PREVENTIVO}}"
            elif text.strip() == "Correctivo":
                node.text = "Correctivo {{TIPO_CORRECTIVO}}"

        if format_type == "tools":
            if "Nueva Instalación" in text:
                node.text = text.replace("Nueva Instalación", "Nueva Instalación {{SERVICIO_NUEVO}}")
            if "Preventivo" in text:
                node.text = text.replace("Preventivo", "Preventivo {{SERVICIO_PREVENTIVO}}")
            if "Correctivo" in text:
                node.text = text.replace("Correctivo", "Correctivo {{SERVICIO_CORRECTIVO}}")

    entries["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone="yes"
    )

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as temp:
        temp_path = Path(temp.name)
    with ZipFile(temp_path, "w", ZIP_DEFLATED) as output:
        for name, data in entries.items():
            output.writestr(name, data)
    temp_path.replace(path)


def main():
    outputs = [
        ("checklist", TEMPLATES / "checklist_diagnostico.docx", TEMPLATES / "checklist_diagnostico_plantilla.docx", prepare_checklist),
        ("quality", TEMPLATES / "pruebas_calidad.docx", TEMPLATES / "pruebas_calidad_plantilla.docx", prepare_quality),
        ("tools", TEMPLATES / "herramientas.docx", TEMPLATES / "herramientas_plantilla.docx", prepare_tools),
    ]
    for format_type, source, destination, builder in outputs:
        builder(source, destination)
        inject_service_fields(destination, format_type)
        print(destination.name)


if __name__ == "__main__":
    main()
